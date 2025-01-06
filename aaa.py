from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

# Read the content of the markdown file
with open('aaa.md', 'r', encoding='utf-8') as md_file:
    md_content = md_file.read()

# Create a new Document
doc = Document()

# Set default font to 游ゴシック
style = doc.styles['Normal']
font = style.font
font.name = '游ゴシック'
r = style.element.rPr.rFonts
r.set(qn('w:eastAsia'), '游ゴシック')

# Initialize counters for headings
heading_counters = [0, 0, 0]

# Function to reset lower level counters
def reset_lower_counters(level):
    for i in range(level, len(heading_counters)):
        heading_counters[i] = 0

# Process each line of the markdown content
table_mode = False
table_data = []
alignments = []
for line in md_content.split('\n'):
    if line.startswith('|'):
        # Process table rows
        table_mode = True
        table_data.append([cell.strip() for cell in line.split('|')[1:-1]])
    elif table_mode and re.match(r'^\s*:\-+\s*|\s*:-+\s*:\s*|\s*-+\s*:\s*$', line):
        # Process table alignment
        alignments = []
        for cell in line.split('|')[1:-1]:
            cell = cell.strip()
            if cell.startswith(':') and cell.endswith(':'):
                alignments.append(WD_ALIGN_PARAGRAPH.CENTER)
            elif cell.startswith(':'):
                alignments.append(WD_ALIGN_PARAGRAPH.LEFT)
            elif cell.endswith(':'):
                alignments.append(WD_ALIGN_PARAGRAPH.RIGHT)
            else:
                alignments.append(WD_ALIGN_PARAGRAPH.LEFT)
    else:
        if table_mode:
            # Create table in docx
            table = doc.add_table(rows=len(table_data) - 1, cols=len(table_data[0]))
            table.style = 'Table Grid'  # Add this line to set the table style with borders
            for row_idx, row_data in enumerate(table_data):
                if row_idx == 1:
                    continue  # Skip the second row
                for col_idx, cell_data in enumerate(row_data):
                    cell = table.cell(row_idx - 1 if row_idx > 1 else row_idx, col_idx)
                    cell.text = cell_data
                    if alignments:
                        cell.paragraphs[0].alignment = alignments[col_idx]
            table_mode = False
            table_data = []

        if line.startswith('#'):
            # Determine the heading level
            level = line.count('#')
            reset_lower_counters(level)
            heading_counters[level - 1] += 1

            # Create the numbering string
            numbering = '-'.join(str(num) for num in heading_counters[:level])

            # Remove the '#' characters and strip leading/trailing whitespace
            heading_text = line.lstrip('#').strip()

            # Format the heading text
            if level == 1:
                heading_text = f"{numbering}、{heading_text}"
                run = doc.add_paragraph().add_run(heading_text)
                run.bold = True
            else:
                heading_text = f"{numbering}、{heading_text}"
                doc.add_paragraph(heading_text)
        elif line.startswith('-'):
            # Add bullet points
            bullet_text = line.lstrip('-').strip()
            doc.add_paragraph(bullet_text, style='List Bullet')
        else:
            # Add normal paragraphs
            doc.add_paragraph(line)

# Set line spacing and paragraph spacing
for paragraph in doc.paragraphs:
    paragraph_format = paragraph.paragraph_format
    paragraph_format.line_spacing = Pt(12)  # Adjust the value as needed
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(0)

# Save the document as a .docx file
doc.save('aaa.docx')

# Read the content of the docx file
doc = Document('aaa.docx')

# Initialize a list to hold the markdown content
md_content = []

# Process each paragraph in the docx file
for para in doc.paragraphs:
    text = para.text.strip()
    if text:
        # Check for heading format
        if re.match(r'^\d+(-\d+)*、', text):
            # Extract the numbering and heading text
            numbering, heading_text = text.split('、', 1)
            level = numbering.count('-') + 1
            md_content.append('#' * level + ' ' + heading_text)
        elif para.style.name == 'List Bullet':
            # Add bullet points
            md_content.append('- ' + text)
        else:
            # Add normal paragraphs
            md_content.append(text)

# Write the markdown content to a new file
with open('bbb.md', 'w', encoding='utf-8') as md_file:
    md_file.write('\n'.join(md_content))
